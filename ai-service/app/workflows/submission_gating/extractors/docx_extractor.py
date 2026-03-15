from __future__ import annotations

import math
from io import BytesIO
from zipfile import ZipFile

from app.workflows.submission_gating.extractors.section_utils import extract_heading_candidates
from app.workflows.submission_gating.models.facts import ExtractedDocument


class DocxExtractor:
    def extract(self, file_bytes: bytes, filename: str | None = None) -> ExtractedDocument:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for DOCX extraction") from exc

        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        raw_text = "\n".join(paragraphs)
        sections = []
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            if style_name.startswith("heading") and paragraph.text.strip():
                sections.append(paragraph.text.strip())
        for section in extract_heading_candidates(paragraphs):
            if section not in sections:
                sections.append(section)

        core_properties = {
            "author": document.core_properties.author or "",
            "title": document.core_properties.title or "",
            "subject": document.core_properties.subject or "",
            "revision": document.core_properties.revision,
        }
        authors = [core_properties["author"]] if core_properties["author"] else []
        page_count = _read_page_count(BytesIO(file_bytes))
        if page_count is None:
            page_count = max(1, math.ceil(len(raw_text.split()) / 500)) if raw_text else 1

        return ExtractedDocument(
            format="docx",
            raw_text=raw_text,
            sections=sections,
            title=core_properties["title"] or (paragraphs[0] if paragraphs else None),
            abstract=_extract_abstract(document, paragraphs),
            authors=authors,
            metadata={},
            table_count=len(document.tables),
            figure_count=sum(1 for shape in document.inline_shapes if shape.type is not None),
            page_count=page_count,
            text_coverage_ratio=min(1.0, len(raw_text) / max(page_count * 500, 1)),
            core_properties=core_properties,
        )

def _extract_abstract(document, paragraphs: list[str]) -> str | None:
    # Check heading-styled paragraphs first
    from docx import Document as _D  # noqa: F401 -- type hint only

    in_abstract = False
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        text = paragraph.text.strip()
        if not text:
            continue
        if style_name.startswith("heading") and text.lower() == "abstract":
            in_abstract = True
            continue
        if in_abstract:
            if style_name.startswith("heading"):
                break
            return text
    # fallback: plain text scan
    for index, paragraph in enumerate(paragraphs):
        if paragraph.lower() == "abstract":
            return " ".join(paragraphs[index + 1 : index + 4]) or None
    return None


def _read_page_count(file_obj: BytesIO) -> int | None:
    try:
        with ZipFile(file_obj) as archive:
            with archive.open("docProps/app.xml") as app_xml:
                content = app_xml.read().decode("utf-8", errors="ignore")
    except Exception:
        return None

    start_tag = "<Pages>"
    end_tag = "</Pages>"
    start = content.find(start_tag)
    end = content.find(end_tag)
    if start == -1 or end == -1:
        return None
    value = content[start + len(start_tag) : end].strip()
    try:
        return int(value)
    except ValueError:
        return None
